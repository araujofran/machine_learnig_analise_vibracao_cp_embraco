from __future__ import annotations

import argparse
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.evaluation.generate_report import load_data


def add_formatted_table(doc, df):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Light Shading Accent 1'
    
    # Headers
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = str(column)
        
    # Data rows
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            if isinstance(val, (float)):
                row_cells[i].text = f"{val:.4f}"
            else:
                row_cells[i].text = str(val)


def generate_word_document(eval_dir: Path, output_docx: Path, img_path: Path):
    df_fault, df_state, df_fp = load_data(eval_dir)

    doc = Document()
    
    # TITULO PRINCIPAL
    title = doc.add_heading('Relatório Analítico Preditivo: Compressor EM2X3125U', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_intro = doc.add_paragraph()
    p_intro.add_run("Este documento executivo detalha a arquitetura front-end e os resultados técnicos do motor preditivo MLOps da Plataforma Embraco.\n").italic = True

    # IMAGEM DO COMPRESSOR
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(5.0))
        img_p = doc.paragraphs[-1]
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Visão Holográfica do Compressor - Gerada por Inteligência Artificial", style='Caption')

    # SEÇÃO 1: ARQUITETURA STREAMLIT
    doc.add_heading('1. Arquitetura do Cockpit (Front-End Streamlit)', level=1)
    
    doc.add_paragraph(
        "Para garantir que os dados de Machine Learning sejam consumíveis "
        "pela liderança técnica de manutenção, construímos uma interface proprietária."
    )
    
    doc.add_heading('Custom CSS (Cyberpunk / Neon Ice)', level=2)
    doc.add_paragraph(
        "Abandonamos a estética branca padrão de data science para aplicar um CSS que remete a um ambiente "
        "de alta-tecnologia industrial. O fundo sideral (#050E17) com fontes 'Orbitron' neon-ice (#00E5FF) causa "
        "impacto e facilita a leitura de alertas visuais (vermelho)."
    )

    doc.add_heading('Plotly Dynamics (Gráficos interativos)', level=2)
    doc.add_paragraph(
        "Todos os gráficos de Z-Score Acústico e Híbrido utilizam a biblioteca Plotly (go.Scatter). "
        "Ela permite o zoom interativo e a plotagem superposta do Threshold (Limiar de Perigo Acústico) como "
        "uma barreira semitransparente, acusando exatamente a janela temporal onde a falha inicial de rolamento cruzou "
        "a banda operacional limite."
    )

    # SEÇÃO 2: RESULTADOS DE VALIDAÇÃO MLOPS
    doc.add_heading('2. Resultados de Validação do Algoritmo', level=1)
    doc.add_paragraph("As métricas a seguir foram consolidadas com base no modelo Isolation Forest combinado com Filtros de Uptime Sensorial.")
    
    # 2.1 UPTIME
    doc.add_heading('A. Falsos Positivos em Regime Não-Produtivo', level=2)
    idle_fp = df_fp[df_fp["operating_state"] == "idle"]["false_positive_rate"].mean() if "idle" in df_fp["operating_state"].values else 0
    startup_fp = df_fp[df_fp["operating_state"] == "startup"]["false_positive_rate"].mean() if "startup" in df_fp["operating_state"].values else 0
    
    p_res1 = doc.add_paragraph()
    if idle_fp < 0.05 and startup_fp < 0.05:
        p_res1.add_run("PERFEITO: ").bold = True
        p_res1.add_run("O filtro de Uptime operou maravilhosamente bem. Alarmes foram suprimidos no Startup e no Repouso.")
    else:
        p_res1.add_run("ATENÇÃO: ").bold = True
        p_res1.add_run(f"O filtro não conteve todos os falsos positivos (Idle: {idle_fp*100:.1f}%).")
        
    add_formatted_table(doc, df_fp)
    
    # 2.2 TAXA DE ACERTO
    doc.add_heading('B. Taxa de Detecção por Falha Mecânica', level=2)
    doc.add_paragraph("Desempenho da detecção das falhas quando o compressor atinge Steady State:")

    for _, row in df_fault.iterrows():
        falha = row["fault_mode"]
        if falha == "healthy":
            continue
            
        rate = float(row["anomaly_rate"])
        rate_pct = rate * 100
        
        if rate_pct >= 90:
            comentario = "Excelente cobertura! O Z-Score híbrido está bruto nessa falha."
        elif rate_pct >= 50:
            comentario = "Aceitável. O alarme ressoou nativamente, porém indica necessidade de baixar o limite de segurança nos próximos deploys."
        else:
            comentario = "Alerta Crítico! Falha sendo largamente mascarada pelo desvio médio limite."
            
        p_falha = doc.add_paragraph(style='List Bullet')
        p_falha.add_run(f"{falha.upper()}: ").bold = True
        p_falha.add_run(f"Acerto em {rate_pct:.1f}% das janelas.\n")
        p_falha.add_run(f"Insights: {comentario}").italic = True

    doc.add_paragraph("\nMatriz Bruta de Avaliação Estatística:")
    add_formatted_table(doc, df_fault)

    doc.add_page_break()
    doc.add_paragraph("Fim do Relatório Corporativo.").italic = True
    
    # Grava documento
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--eval_dir", type=str, default="data/processed")
    parser.add_argument("--output_docx", type=str, default="docs/Relatorio_Apresentacao_Dashboard.docx")
    parser.add_argument("--img_path", type=str, default="assets/compressor_hero.png")
    args = parser.parse_args()

    eval_dir = Path(args.eval_dir)
    out_file = Path(args.output_docx)
    img_path = Path(args.img_path)

    generate_word_document(eval_dir, out_file, img_path)
    print(f"\n[OK] Documento de Apresentação Word gravado em: {out_file}\n")


if __name__ == "__main__":
    main()
